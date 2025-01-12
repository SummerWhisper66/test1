import os
from docx import Document
from docx.shared import Cm
from PIL import Image
from PyQt5.QtWidgets import QApplication, QWidget, QVBoxLayout, QHBoxLayout, QPushButton, QLineEdit, QLabel, \
    QFileDialog, QProgressBar, QSpinBox, QFormLayout, QDoubleSpinBox
from PyQt5.QtCore import QThread, pyqtSignal, Qt
from PyQt5.QtGui import QPixmap, QIcon, QPalette, QBrush, QPainter


class Worker(QThread):
    progress_changed = pyqtSignal(int)

    def __init__(self, root_folder, output_folder, rows, cols, new_width, new_height, image_width, image_height):
        super().__init__()
        self.root_folder = root_folder
        self.output_folder = output_folder
        self.rows = rows
        self.cols = cols
        self.new_width = new_width
        self.new_height = new_height
        self.image_width = image_width
        self.image_height = image_height

    def run(self):
        # 获取所有图片文件数量
        total_images = sum(
            [len([f for f in os.listdir(os.path.join(self.root_folder, d)) if f.endswith(('.png', '.jpg', '.jpeg'))])
             for d in os.listdir(self.root_folder) if os.path.isdir(os.path.join(self.root_folder, d))])

        current_image = 0

        # 遍历文件夹
        for dirpath, dirnames, filenames in os.walk(self.root_folder):
            for dirname in dirnames:
                # 创建一个空文档
                document = Document()
                # 创建一个空表格，表格为自定义的行和列
                table = document.add_table(rows=self.rows, cols=self.cols)

                subfolder_path = os.path.join(dirpath, dirname)
                image_files = [f for f in os.listdir(subfolder_path) if f.endswith(('.png', '.jpg', '.jpeg'))]

                row = 0
                col = 0
                for image_file in image_files:
                    image_path = os.path.join(subfolder_path, image_file)

                    # 修改图片尺寸
                    image = Image.open(image_path)
                    width, height = image.size
                    if width > self.new_width:
                        height = int(self.new_width * height / width)
                        width = self.new_width
                    if height > self.new_height:
                        width = int(self.new_height * width / height)
                        height = self.new_height

                    image.resize((width, height)).save(image_path)

                    # 将图片插入表格
                    cell = table.cell(row, col)
                    paragraph = cell.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=Cm(self.image_width), height=Cm(self.image_height))

                    col += 1
                    if col == self.cols:
                        col = 0
                        row += 1

                    current_image += 1
                    progress = int((current_image / total_images) * 100)
                    self.progress_changed.emit(progress)

                # 保存文档
                docx_path = os.path.join(self.output_folder, f'{dirname}.docx')
                document.save(docx_path)


class App(QWidget):
    def __init__(self):
        super().__init__()

        self.setWindowTitle('图片插入Word表格工具')
        self.resize(550, 400)
        self.setWindowIcon(QIcon('./icon/00002.png'))  # 设置图标

        # 加载背景图片
        self.background_pixmap = QPixmap("./Pictures/00031-1149160498.png")

        '''
        # 设置自动填充背景
        # self.setAutoFillBackground(True)
        
        # 设置背景色为指定图片
        # p = self.palette()
        # p.setBrush(self.backgroundRole(), QBrush(QPixmap("F:/Python_Files/QC_Check/Pictures/00031-1149160498.png")))
        # self.setPalette(p)

        # 设置背景图像
        # self.setStyleSheet(
        #     "QWidget { "
        #     "background-image: url('F:/Python_Files/QC_Check/Pictures/00031-1149160498.png'); "
        #     "background-repeat: no-repeat; "
        #     "background-size: cover; "
        #     "background-position: center; "
        #     "}")
        '''

        # UI组件
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()

        # 文件夹路径输入部分
        self.folder_input_layout = QHBoxLayout()
        self.folder_input_label = QLabel('选择图片文件夹:')
        self.folder_input_line = QLineEdit(self)
        self.folder_input_button = QPushButton('浏览', self)
        self.folder_input_button.clicked.connect(self.select_folder)
        self.folder_input_layout.addWidget(self.folder_input_label)
        self.folder_input_layout.addWidget(self.folder_input_line)
        self.folder_input_layout.addWidget(self.folder_input_button)

        # 输出文件夹路径输入部分
        self.output_layout = QHBoxLayout()
        self.output_label = QLabel('选择输出文件夹:')
        self.output_line = QLineEdit(self)
        self.output_button = QPushButton('浏览', self)
        self.output_button.clicked.connect(self.select_output_folder)
        self.output_layout.addWidget(self.output_label)
        self.output_layout.addWidget(self.output_line)
        self.output_layout.addWidget(self.output_button)

        # 表格行列输入
        self.table_layout = QFormLayout()
        self.rows_input = QSpinBox(self)
        self.rows_input.setValue(60)  # 默认行数
        self.cols_input = QSpinBox(self)
        self.cols_input.setValue(2)   # 默认列数
        self.table_layout.addRow('表格行数:', self.rows_input)
        self.table_layout.addRow('表格列数:', self.cols_input)

        # 图片尺寸输入
        self.size_layout = QFormLayout()
        self.new_width_input = QSpinBox(self)
        self.new_width_input.setRange(1, 10000)  # 设置更大的范围（例如：1到10000）
        self.new_width_input.setValue(2000)  # 默认图片宽度
        self.new_height_input = QSpinBox(self)
        self.new_height_input.setRange(1, 10000)  # 设置更大的范围（例如：1到10000）
        self.new_height_input.setValue(1500)  # 默认图片高度
        self.size_layout.addRow('图片目标宽度:', self.new_width_input)
        self.size_layout.addRow('图片目标高度:', self.new_height_input)

        # 单元格图片大小输入
        self.cell_size_layout = QFormLayout()
        self.cell_width_input = QDoubleSpinBox(self)
        self.cell_width_input.setValue(7.51)  # 默认单元格图片宽度
        self.cell_height_input = QDoubleSpinBox(self)
        self.cell_height_input.setValue(5.64)  # 默认单元格图片高度
        self.cell_size_layout.addRow('插入单元格图片宽度 (cm):', self.cell_width_input)
        self.cell_size_layout.addRow('插入单元格图片高度 (cm):', self.cell_height_input)

        # 进度条
        self.progress_bar = QProgressBar(self)
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)

        # 开始按钮
        self.start_button = QPushButton('开始处理', self)
        self.start_button.clicked.connect(self.start_processing)

        # 设置鼠标悬浮提示
        self.start_button.setToolTip("注意该程序只会同比例压缩图片文件，填入的图片像素比例若与原图片比例不同，程序会按照原图片比例来压缩")

        # 添加版权信息
        footer_layout = QHBoxLayout()
        self.footer_label = QLabel("Copyright:MiemieY|2025.01.12")  # 创建版权信息
        self.footer_label.setAlignment(Qt.AlignLeft | Qt.AlignBottom)
        footer_layout.addWidget(self.footer_label)

        # 将UI元素添加到主布局
        layout.addLayout(self.folder_input_layout)
        layout.addLayout(self.output_layout)
        layout.addLayout(self.table_layout)
        layout.addLayout(self.size_layout)
        layout.addLayout(self.cell_size_layout)
        layout.addWidget(self.progress_bar)
        layout.addWidget(self.start_button)
        layout.addLayout(footer_layout)

        self.setLayout(layout)

        # 居中显示UI
        self.center()

    def center(self):
        frame_geometry = self.frameGeometry()
        screen = QApplication.desktop().screenNumber(QApplication.desktop().cursor().pos())
        center_point = QApplication.desktop().screenGeometry(screen).center()
        frame_geometry.moveCenter(center_point)
        self.move(frame_geometry.topLeft())

    def select_folder(self):
        folder = QFileDialog.getExistingDirectory(self, '选择图片文件夹')
        if folder:
            self.folder_input_line.setText(folder)

    def select_output_folder(self):
        folder = QFileDialog.getExistingDirectory(self, '选择输出文件夹')
        if folder:
            self.output_line.setText(folder)

    def start_processing(self):
        root_folder = self.folder_input_line.text()
        output_folder = self.output_line.text()
        rows = self.rows_input.value()
        cols = self.cols_input.value()
        new_width = self.new_width_input.value()
        new_height = self.new_height_input.value()
        image_width = self.cell_width_input.value()
        image_height = self.cell_height_input.value()

        if not root_folder or not output_folder:
            return

        self.worker = Worker(root_folder, output_folder, rows, cols, new_width, new_height, image_width, image_height)
        self.worker.progress_changed.connect(self.update_progress)
        self.worker.start()

    def update_progress(self, progress):
        self.progress_bar.setValue(progress)


    def paintEvent(self, event):
        painter = QPainter(self)
        window_size = self.size()

        # 将背景图像调整为窗口大小
        scaled_pixmap = self.background_pixmap.scaled(window_size, Qt.KeepAspectRatioByExpanding,
                                                      Qt.SmoothTransformation)

        # 绘制背景图像
        painter.drawPixmap(0, 0, scaled_pixmap)
        painter.end()

if __name__ == '__main__':
    app = QApplication([])
    window = App()
    window.show()
    app.exec_()
