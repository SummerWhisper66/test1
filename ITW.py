import os
from docx import Document
from docx.shared import Cm
from PIL import Image
from PyQt5.QtWidgets import QApplication, QWidget, QVBoxLayout, QHBoxLayout, QPushButton, QLineEdit, QLabel, \
    QFileDialog, QProgressBar, QGraphicsOpacityEffect, QFrame
from PyQt5.QtCore import QThread, pyqtSignal, Qt
from PyQt5.QtGui import QPixmap, QIcon


class Worker(QThread):
    progress_changed = pyqtSignal(int)

    def __init__(self, root_folder, output_folder):
        super().__init__()
        self.root_folder = root_folder
        self.output_folder = output_folder

    def run(self):
        # 获取所有图片文件数量
        total_images = sum(
            [len([f for f in os.listdir(os.path.join(self.root_folder, d)) if f.endswith(('.png', '.jpg', '.jpeg'))])
             for d in os.listdir(self.root_folder) if os.path.isdir(os.path.join(self.root_folder, d))])

        current_image = 0

        # 遍历文件夹
        for dirpath, dirnames, filenames in os.walk(self.root_folder):
            for dirname in dirnames:
                # 创建一个空文档
                document = Document()
                table = document.add_table(rows=50, cols=2)

                subfolder_path = os.path.join(dirpath, dirname)
                image_files = [f for f in os.listdir(subfolder_path) if f.endswith(('.png', '.jpg', '.jpeg'))]

                row = 0
                col = 0
                for image_file in image_files:
                    image_path = os.path.join(subfolder_path, image_file)

                    # 修改图片尺寸
                    image = Image.open(image_path)
                    width, height = image.size
                    new_width, new_height = 2000, 1500
                    if width > new_width:
                        height = int(new_width * height / width)
                        width = new_width
                    if height > new_height:
                        width = int(new_height * width / height)
                        height = new_height

                    image.resize((new_width, new_height)).save(image_path)

                    # 将图片插入表格
                    cell = table.cell(row, col)
                    paragraph = cell.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=Cm(7.51), height=Cm(5.64))

                    col += 1
                    if col == 2:
                        col = 0
                        row += 1

                    current_image += 1
                    progress = int((current_image / total_images) * 100)
                    self.progress_changed.emit(progress)

                # 保存文档
                docx_path = os.path.join(self.output_folder, f'{dirname}.docx')
                document.save(docx_path)


class App(QWidget):
    def __init__(self):
        super().__init__()

        self.setWindowTitle('图片插入Word表格工具')
        self.setGeometry(300, 200, 600, 250)
        self.setWindowIcon(QIcon('F:\\Python_Files\\QC_Check\\icon\\00002.png'))

        self.setAutoFillBackground(True)
        p = self.palette()
        p.setBrush(self.backgroundRole(), Qt.white)
        self.setPalette(p)

        self.setStyleSheet(
            "QWidget { background-image: url(F:\\Python_Files\\QC_Check\\Pictures\\00006-1155431150.png); background-repeat: no-repeat; background-position: center; }")

        # UI组件
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()

        # 输入文件夹部分
        self.folder_input_layout = QHBoxLayout()
        self.folder_input_label = QLabel('选择图片文件夹:')
        self.folder_input_line = QLineEdit(self)
        self.folder_input_button = QPushButton('浏览', self)
        self.folder_input_button.clicked.connect(self.select_folder)
        self.folder_input_layout.addWidget(self.folder_input_label)
        self.folder_input_layout.addWidget(self.folder_input_line)
        self.folder_input_layout.addWidget(self.folder_input_button)

        # 输出文件夹部分
        self.output_layout = QHBoxLayout()
        self.output_label = QLabel('选择输出文件夹:')
        self.output_line = QLineEdit(self)
        self.output_button = QPushButton('浏览', self)
        self.output_button.clicked.connect(self.select_output_folder)
        self.output_layout.addWidget(self.output_label)
        self.output_layout.addWidget(self.output_line)
        self.output_layout.addWidget(self.output_button)

        # 进度条
        self.progress_bar = QProgressBar(self)
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)

        # 开始按钮
        self.start_button = QPushButton('开始处理', self)
        self.start_button.clicked.connect(self.start_processing)

        # 版权信息
        self.copyright_label = QLabel('Copyright: MiemieY', self)
        self.copyright_label.setAlignment(Qt.AlignRight | Qt.AlignBottom)  # 右下角显示
        self.copyright_label.setStyleSheet('color: white; font-size: 12px;')

        # 将UI元素添加到主布局
        layout.addLayout(self.folder_input_layout)
        layout.addLayout(self.output_layout)
        layout.addWidget(self.progress_bar)
        layout.addWidget(self.start_button)
        layout.addWidget(self.copyright_label)  # 版权信息放在最下面

        self.setLayout(layout)

        # 居中显示UI
        self.center()

    def center(self):
        frame_geometry = self.frameGeometry()
        screen = QApplication.desktop().screenNumber(QApplication.desktop().cursor().pos())
        center_point = QApplication.desktop().screenGeometry(screen).center()
        frame_geometry.moveCenter(center_point)
        self.move(frame_geometry.topLeft())

    def select_folder(self):
        folder = QFileDialog.getExistingDirectory(self, '选择图片文件夹')
        if folder:
            self.folder_input_line.setText(folder)

    def select_output_folder(self):
        folder = QFileDialog.getExistingDirectory(self, '选择输出文件夹')
        if folder:
            self.output_line.setText(folder)

    def start_processing(self):
        root_folder = self.folder_input_line.text()
        output_folder = self.output_line.text()

        if not root_folder or not output_folder:
            return

        self.worker = Worker(root_folder, output_folder)
        self.worker.progress_changed.connect(self.update_progress)
        self.worker.start()

    def update_progress(self, progress):
        self.progress_bar.setValue(progress)


if __name__ == '__main__':
    app = QApplication([])
    window = App()
    window.show()
    app.exec_()
